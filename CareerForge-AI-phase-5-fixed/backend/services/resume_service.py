"""Resume parsing and deterministic analysis for CareerForge Phase 5.

This module deliberately keeps resume-specific logic out of main.py.  The
fallback is content based: it never invents candidate facts and remains
available when Gemini is not configured.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from pypdf import PdfReader


MAX_UPLOAD_BYTES = 8 * 1024 * 1024

SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "contact": ("contact",),
    "summary": ("summary", "professional summary", "profile", "about me"),
    "objective": ("objective", "career objective"),
    "skills": ("skills", "technical skills", "technologies", "tech stack"),
    "experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "internship",
    ),
    "projects": ("projects", "personal projects", "academic projects", "selected projects"),
    "education": ("education", "academic background"),
    "certifications": ("certifications", "certificates", "licenses"),
    "achievements": ("achievements", "awards", "honors"),
    "publications": ("publications", "research", "papers"),
    "links": ("links", "profiles", "online profiles", "portfolio"),
}

COMMON_ROLE_KEYWORDS: dict[str, tuple[str, ...]] = {
    "software engineer": (
        "software development", "data structures", "algorithms", "testing",
        "api", "git", "cloud", "database",
    ),
    "backend": (
        "api", "rest", "database", "sql", "testing", "authentication",
        "docker", "cloud", "microservices",
    ),
    "java": ("java", "spring", "spring boot", "rest", "sql", "junit", "maven"),
    "frontend": (
        "javascript", "typescript", "react", "html", "css", "responsive",
        "testing", "accessibility",
    ),
    "full stack": (
        "javascript", "typescript", "react", "api", "database", "sql",
        "authentication", "deployment",
    ),
    "data analyst": (
        "sql", "python", "excel", "tableau", "power bi", "statistics",
        "data visualization", "analytics",
    ),
    "devops": (
        "linux", "docker", "kubernetes", "ci/cd", "aws", "azure",
        "terraform", "monitoring",
    ),
    "machine learning": (
        "python", "pandas", "numpy", "scikit-learn", "model", "statistics",
        "evaluation", "deployment",
    ),
    "product manager": (
        "roadmap", "user research", "requirements", "analytics", "stakeholder",
        "product strategy",
    ),
}

ACTION_VERBS = (
    "built", "developed", "designed", "implemented", "created", "led",
    "optimized", "automated", "deployed", "analyzed", "launched", "delivered",
    "architected", "refactored", "tested", "configured", "migrated",
)
TECHNOLOGY_TERMS = (
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", "golang",
    "rust", "ruby", "php", "swift", "kotlin", "sql", "nosql", "html", "css",
    "react", "next.js", "node.js", "node", "angular", "vue", "svelte", "spring",
    "spring boot", "django", "flask", "fastapi", "express", "graphql", "rest",
    "api", "postgresql", "postgres", "mysql", "mongodb", "redis", "sqlite",
    "docker", "docker compose", "kubernetes", "aws", "azure", "gcp", "terraform",
    "git", "github", "gitlab", "jenkins", "linux", "pandas", "numpy",
    "scikit-learn", "tensorflow", "pytorch", "tableau", "power bi", "excel",
)
WEAK_STARTS = (
    "worked on", "responsible for", "helped with", "worked with",
    "was responsible", "assisted with", "involved in",
)
DATE_PATTERN = re.compile(
    r"\b(?:19|20)\d{2}\b|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)"
    r"[a-z]*\.?\s+(?:19|20)?\d{2}\b",
    re.IGNORECASE,
)
EMAIL_PATTERN = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)(?!\d)")


def _normalise_space(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def normalise_for_comparison(value: str) -> str:
    """Normalise case, whitespace, and harmless punctuation for comparisons."""
    value = _normalise_space(value).casefold()
    return re.sub(r"[“”‘’`]", "'", value)


def _protected_facts(text: str) -> set[str]:
    """Return facts that a rewrite must not introduce from nowhere."""
    value = text or ""
    facts = {
        match.group(0).casefold()
        for pattern in (
            EMAIL_PATTERN,
            re.compile(r"https?://[^\s<>()]+", re.I),
            re.compile(r"\b(?:19|20)\d{2}\b"),
            re.compile(r"\b\d+(?:\.\d+)?%?\b"),
        )
        for match in pattern.finditer(value)
    }
    # Technology names are a conservative, high-value subset of factual claims.
    low = value.casefold()
    facts.update(term for term in TECHNOLOGY_TERMS if _contains_keyword(value, term))
    return facts


def _new_protected_facts(original: str, suggested: str) -> set[str]:
    original_facts = _protected_facts(original)
    return _protected_facts(suggested) - original_facts


def _new_proper_terms(original: str, suggested: str) -> set[str]:
    """Catch likely new names/titles while allowing normal sentence starts."""
    term_pattern = r"[A-Za-z][A-Za-z0-9+#-]*(?:\.[A-Za-z0-9+#-]+)*"
    original_words = {word.casefold() for word in re.findall(term_pattern, original or "") if len(word) > 2}
    new_terms: set[str] = set()
    for line in (suggested or "").splitlines():
        words = re.findall(rf"\b[A-Z][A-Za-z0-9+#-]*(?:\.[A-Za-z0-9+#-]+)*\b", line)
        first_word = True
        for word in words:
            if first_word:
                first_word = False
                continue
            if word.casefold() not in original_words:
                new_terms.add(word)
    return new_terms


def validate_improvement(
    resume_text: str,
    current_text: str,
    suggested_text: str,
) -> list[str]:
    """Validate a proposed replacement without pretending to prove semantics."""
    current = (current_text or "").strip()
    suggested = (suggested_text or "").strip()
    errors: list[str] = []
    if not suggested:
        errors.append("The suggested text cannot be empty.")
    if current and not current.startswith("[") and current not in (resume_text or ""):
        errors.append("The current text is no longer present in the working resume.")
    if current and normalise_for_comparison(current) == normalise_for_comparison(suggested):
        errors.append("The suggestion must make a meaningful wording change.")
    new_facts = _new_protected_facts(resume_text, suggested)
    new_terms = _new_proper_terms(resume_text, suggested)
    new_facts.update(term.casefold() for term in new_terms)
    if new_facts:
        errors.append(
            "The suggestion introduces unsupported factual details: "
            + ", ".join(sorted(new_facts))
            + "."
        )
    return errors


def _tokens(value: str) -> set[str]:
    return {
        token for token in re.findall(r"[a-z0-9+#.]+", (value or "").lower())
        if len(token) > 1
    }


def _contains_keyword(text: str, keyword: str) -> bool:
    value = (text or "").lower()
    if re.search(r"[^a-z0-9]", keyword):
        return keyword.lower() in value
    return bool(re.search(rf"(?<![a-z0-9+#.]){re.escape(keyword.lower())}(?![a-z0-9+#.])", value))


def _section_heading(line: str) -> str | None:
    candidate = re.sub(r"^[\s•·▪▸\-–—\d.)]+", "", line.strip()).rstrip(":").strip().lower()
    candidate = re.sub(r"\s+", " ", candidate)
    for section, aliases in SECTION_ALIASES.items():
        if candidate in aliases:
            return section
    return None


def detect_sections(text: str) -> dict[str, bool]:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    result = {name: False for name in SECTION_ALIASES}
    for line in lines:
        heading = _section_heading(line)
        if heading:
            result[heading] = True
    # Contact details often have no "Contact" heading.
    result["contact"] = bool(EMAIL_PATTERN.search(text or "") or PHONE_PATTERN.search(text or ""))
    # A common compact resume format omits a heading but still has these signals.
    low = (text or "").lower()
    result["links"] = result["links"] or bool(re.search(r"https?://|linkedin\.com|github\.com", low))
    return result


def extract_resume(filename: str, data: bytes) -> str:
    """Extract text without executing or persisting the uploaded file."""
    lower = (filename or "").lower().strip()
    try:
        if lower.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(data))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if lower.endswith(".docx"):
            doc = Document(io.BytesIO(data))
            chunks: list[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    chunks.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        chunks.append(row_text)
            return "\n".join(chunks).strip()
        if lower.endswith(".txt"):
            text = data.decode("utf-8-sig", errors="replace").strip()
            if text.count("\ufffd") > max(5, len(text) // 80):
                raise ValueError("The text file does not appear to be readable UTF-8 text.")
            return text
    except Exception as exc:
        # Callers turn this into a safe user-facing error; implementation details
        # and parser stack traces never cross the API boundary.
        raise ValueError("The file could not be read. Try exporting it again as PDF, DOCX, or TXT.") from exc
    raise ValueError("Supported formats: PDF, DOCX, TXT.")


def _section_content(text: str, target: str) -> str:
    lines = (text or "").splitlines()
    start = None
    for index, line in enumerate(lines):
        if _section_heading(line) == target:
            start = index + 1
            break
    if start is None:
        return ""
    end = len(lines)
    for index in range(start, len(lines)):
        if _section_heading(lines[index]):
            end = index
            break
    return "\n".join(lines[start:end]).strip()


def _content_lines(text: str, section: str | None = None) -> list[str]:
    value = _section_content(text, section) if section else text
    return [line.strip() for line in value.splitlines() if line.strip()]


def _bullet_lines(text: str, section: str | None = None) -> list[str]:
    return [
        line for line in _content_lines(text, section)
        if line.startswith(("-", "•", "*", "·", "–", "—")) or len(line.split()) >= 8
    ]


def _role_keywords(target_role: str) -> list[str]:
    role = _normalise_space(target_role).lower()
    selected: list[str] = []
    for key, keywords in COMMON_ROLE_KEYWORDS.items():
        if key in role:
            selected.extend(keywords)
    if not selected:
        selected.extend(
            phrase for phrase in re.findall(r"[a-z][a-z0-9+#.-]{2,}(?:\s+[a-z][a-z0-9+#.-]{2,})?", role)
            if phrase not in {"developer", "engineer", "professional", "specialist", "analyst"}
        )
    # Role words are useful evidence, while generic words are not useful gaps.
    selected.extend(token for token in _tokens(role) if token not in {
        "developer", "engineer", "professional", "specialist", "analyst", "manager",
    })
    deduped: list[str] = []
    for keyword in selected:
        keyword = _normalise_space(keyword)
        if keyword and keyword not in deduped:
            deduped.append(keyword)
    return deduped[:18]


def _format_score(value: float) -> float:
    return round(max(0.0, min(10.0, value)), 1)


def _category_scores(text: str, sections: dict[str, bool], target_role: str) -> dict[str, float]:
    words = text.split()
    lines = _content_lines(text)
    bullets = _bullet_lines(text)
    low = text.lower()
    keywords = _role_keywords(target_role)
    keyword_hits = sum(_contains_keyword(text, keyword) for keyword in keywords)
    action_count = sum(
        bool(re.search(rf"\b{re.escape(verb)}\b", low)) for verb in ACTION_VERBS
    )
    quantified = len(re.findall(
        r"\b\d+(?:\.\d+)?%|\b\d+\+|\b\d+\s*(?:users|customers|projects|requests|ms|seconds|days)\b",
        low,
    ))
    heading_score = sum(sections.values()) / len(sections) * 10
    ats = 4.4
    ats += 1.2 if sections["contact"] else 0
    ats += 1.0 if sum(sections.values()) >= 5 else 0
    ats += 0.7 if not re.search(r"[│║]{2,}|[^\S\r\n]{8,}", text) else -0.8
    ats -= 0.5 if len([line for line in lines if len(line) > 180]) > 2 else 0
    content = 3.6 + min(2.0, len(words) / 260) + min(1.3, action_count * 0.22) + min(1.1, quantified * 0.28)
    role = 3.0 if not target_role else 3.1 + (5.8 * keyword_hits / max(1, len(keywords)))
    skills = 3.0 + (4.8 if sections["skills"] else 0) + min(2.0, len(_section_content(text, "skills").split(",")) * 0.15)
    experience = 3.0 + (4.2 if sections["experience"] else 0) + min(2.0, action_count * 0.25) + min(0.8, quantified * 0.15)
    projects = 3.0 + (4.2 if sections["projects"] else 0) + (1.0 if "github" in low or "live demo" in low else 0)
    achievements = 2.5 + (3.0 if sections["achievements"] else 0) + min(4.0, quantified * 0.8)
    formatting = 4.0 + min(2.5, heading_score / 2.5) - (1.1 if re.search(r"[│║]{2,}", text) else 0)
    keyword_score = 3.0 if not keywords else 2.5 + 7.0 * keyword_hits / len(keywords)
    readability = 4.8 + (1.0 if 250 <= len(words) <= 900 else 0) + (1.0 if len(lines) >= 8 else 0)
    readability -= 1.1 if len(words) > 1100 else 0
    return {
        "ats_compatibility": _format_score(ats),
        "content_quality": _format_score(content),
        "role_alignment": _format_score(role),
        "skills": _format_score(skills),
        "experience": _format_score(experience),
        "projects": _format_score(projects),
        "achievements": _format_score(achievements),
        "formatting_structure": _format_score(formatting),
        "keyword_coverage": _format_score(keyword_score),
        "readability": _format_score(readability),
    }


def _keyword_analysis(text: str, target_role: str) -> dict[str, Any]:
    keywords = _role_keywords(target_role)
    strong = [keyword for keyword in keywords if _contains_keyword(text, keyword)]
    missing = [keyword for keyword in keywords if keyword not in strong]
    counts: dict[str, int] = {}
    for token in re.findall(r"[a-z][a-z0-9+#.-]{2,}", text.lower()):
        counts[token] = counts.get(token, 0) + 1
    overused = [
        token for token, count in sorted(counts.items(), key=lambda item: (-item[1], item[0]))
        if count >= 5 and token not in {"experience", "project", "skills", "development"}
    ][:8]
    return {
        "strong_keywords": strong[:12],
        "missing_keywords": missing[:12],
        "overused_keywords": overused,
        "suggested_keywords": missing[:10],
        "warning": "Only add a suggested keyword when you genuinely have experience with it.",
    }


def _ats_analysis(text: str, sections: dict[str, bool], target_role: str) -> dict[str, Any]:
    lines = _content_lines(text)
    weak = [
        line for line in lines
        if any(line.lower().startswith(prefix) for prefix in WEAK_STARTS)
    ]
    issues: list[str] = []
    if not sections["contact"]:
        issues.append("Missing clearly detectable contact information such as an email or phone number.")
    if not sections["summary"] and not sections["objective"]:
        issues.append("No Summary or Objective heading was detected.")
    if sum(sections.values()) < 5:
        issues.append("Several standard sections are not clearly labeled.")
    if re.search(r"[│║]{2,}|[^\S\r\n]{8,}", text):
        issues.append("Spacing or column-like symbols may be difficult for some parsers to interpret.")
    if not DATE_PATTERN.search(text):
        issues.append("No clear year or month-year dates were detected.")
    if weak:
        issues.append(f"{len(weak)} bullet(s) begin with generic responsibility language.")
    if len([line for line in lines if len(line) > 180]) > 2:
        issues.append("Some sections contain unusually long lines; shorter bullets are easier to scan.")
    if target_role and not _role_keywords(target_role):
        issues.append("The target role did not provide enough recognizable terms for a precise keyword comparison.")
    return {
        "label": "ATS-style compatibility analysis",
        "issues": issues[:10],
        "missing_keywords": _keyword_analysis(text, target_role)["missing_keywords"],
        "weak_bullets": weak[:8],
        "formatting_flags": {
            "excessive_formatting_detected": bool(re.search(r"[│║]{2,}", text)),
            "column_like_spacing_detected": bool(re.search(r"[^\S\r\n]{8,}", text)),
            "unclear_section_headings": sum(sections.values()) < 5,
            "unclear_dates": not bool(DATE_PATTERN.search(text)),
            "possible_keyword_stuffing": len(_keyword_analysis(text, target_role)["overused_keywords"]) >= 3,
        },
        "note": "This is an ATS-style analysis based on readable text, not a reproduction of any specific ATS vendor.",
    }


def _role_alignment(text: str, target_role: str) -> dict[str, Any]:
    keywords = _role_keywords(target_role)
    evidence = [
        {
            "keyword": keyword,
            "status": "present" if _contains_keyword(text, keyword) else "missing",
            "label": "✓" if _contains_keyword(text, keyword) else "✗",
        }
        for keyword in keywords
    ]
    present = [item["keyword"] for item in evidence if item["status"] == "present"]
    missing = [item["keyword"] for item in evidence if item["status"] == "missing"]
    score = _format_score(3.0 if not target_role else 3.0 + 7 * len(present) / max(1, len(evidence)))
    return {
        "target_role": target_role or "Not specified",
        "evidence": evidence,
        "score": score,
        "matches": present[:10],
        "missing": missing[:10],
        "emphasize": present[:5] or ["Add role-relevant evidence that is already true of your experience."],
        "add_only_if_true": missing[:8],
        "guidance": "Missing terms are prompts to check your experience, not claims to copy into the resume.",
    }


def _project_analysis(text: str) -> list[dict[str, Any]]:
    content = _section_content(text, "projects")
    if not content:
        return []
    lines = _content_lines(content)
    if not lines:
        return []
    # Keep every project signal. When a resume has no obvious project headings,
    # treat the section as one entry rather than inventing project boundaries.
    candidates: list[str] = []
    for line in lines:
        if (
            len(line.split()) <= 9
            and not line.startswith(("-", "•", "*", "·"))
            and not DATE_PATTERN.search(line)
        ):
            candidates.append(line.rstrip(":"))
    if not candidates:
        candidates = ["Projects section"]
    projects: list[dict[str, Any]] = []
    for index, name in enumerate(candidates):
        projects.append({
            "name": name,
            "technologies": [
                token for token in re.findall(r"\b[A-Z][A-Za-z0-9+#.-]{1,}\b", content)
                if token.lower() not in {"Project", "Built", "Created"}
            ][:12],
            "purpose": _normalise_space(content)[:240],
            "complexity": "Evidence is limited to the uploaded resume text.",
            "technical_depth": "Review the implementation details, architecture, testing, and deployment evidence.",
            "measurable_impact": (
                "Provided in the resume." if re.search(r"\b\d+(?:\.\d+)?%|\b\d+\+", content)
                else "No measurable impact was detected; add one only if you can verify it."
            ),
            "github_link": next((url for url in re.findall(r"https?://\S+", content) if "github" in url.lower()), ""),
            "live_demo": next((url for url in re.findall(r"https?://\S+", content) if "github" not in url.lower()), ""),
            "quality_score": _format_score(4.0 + (2 if len(content.split()) > 35 else 0) + (2 if re.search(r"\b(?:api|database|deployed|test)\b", content, re.I) else 0)),
            "how_to_strengthen": "Name the problem, technical decisions, your contribution, and a verified outcome.",
        })
    return projects


def _experience_analysis(text: str, target_role: str) -> list[dict[str, Any]]:
    content = _section_content(text, "experience")
    if not content:
        return []
    lines = _content_lines(content)
    entries = [line for line in lines if DATE_PATTERN.search(line) or len(line.split()) <= 10]
    if not entries:
        entries = ["Experience section"]
    role_words = set(_tokens(target_role))
    result = []
    for entry in entries[:8]:
        result.append({
            "entry": entry,
            "job_title": entry,
            "organization": "Not separately detected",
            "dates": "Detected in resume." if DATE_PATTERN.search(content) else "Dates not clearly detected.",
            "responsibilities": "Review the bullets under this entry for task descriptions.",
            "technologies": [token for token in re.findall(r"\b[A-Z][A-Za-z0-9+#.-]{1,}\b", content)][:10],
            "achievements": "Verified achievements or metrics are not separately labeled.",
            "measurable_impact": "Provided in the resume." if re.search(r"\b\d+(?:\.\d+)?%|\b\d+\+", content) else "Add a measurable result if you can verify one.",
            "relevance": "Relevant evidence detected." if role_words & _tokens(content) else "Make relevance to the target role more explicit if truthful.",
            "writing_guidance": "Distinguish what you were responsible for from what changed because of your work.",
        })
    return result


def _bullet_suggestions(text: str) -> list[dict[str, str]]:
    suggestions: list[dict[str, str]] = []
    for section in ("experience", "projects"):
        for line in _bullet_lines(text, section):
            clean = re.sub(r"^[•·▪▸*\-–—]\s*", "", line).strip()
            lower = clean.lower()
            prefix = next((item for item in WEAK_STARTS if lower.startswith(item)), None)
            if not prefix:
                continue
            remainder = clean[len(prefix):].lstrip(" :,-")
            replacements = {
                "worked on": "Developed",
                "worked with": "Used",
                "responsible for": "Managed",
                "was responsible for": "Managed",
                "helped with": "Supported",
                "assisted with": "Supported",
                "involved in": "Contributed to",
            }
            suggested = next(
                (
                    f"{replacement} {clean[len(prefix):].lstrip(' :,-')}"
                    for weak, replacement in replacements.items()
                    if lower.startswith(weak)
                ),
                f"Contributed to {remainder[:1].lower() + remainder[1:]}" if remainder
                else "Describe the specific contribution",
            )
            suggestions.append({
                "section": section.title(),
                "current": line,
                "suggested": suggested.rstrip(".") + ".",
                "why": "Starts with a clear action and leaves room for factual technology and outcome evidence.",
            })
    return suggestions[:6]


def _deterministic_rewrite(current: str, section: str) -> str:
    """Make a conservative wording or formatting improvement using existing text."""
    value = current.strip()
    bullet = re.match(r"^([-•*·–—])\s+", value)
    bullet_prefix = bullet.group(1) if bullet else ""
    rewrite_value = re.sub(r"^[-•*·–—]\s+", "", value)
    replacements = (
        (r"^\s*(?:worked on|was responsible for|responsible for)\b", "Managed"),
        (r"^\s*worked with\b", "Used"),
        (r"^\s*(?:helped with|assisted with)\b", "Supported"),
        (r"^\s*involved in\b", "Contributed to"),
        (r"^\bparticipated in\b", "Contributed to"),
        (r"\butilized\b", "Used"),
        (r"\bin order to\b", "to"),
        (r"\bwas able to\b", "could"),
        (r"\bcreated\b", "Developed"),
        (r"\bbuilt\b", "Developed"),
    )
    for pattern, replacement in replacements:
        rewritten, count = re.subn(pattern, replacement, rewrite_value, count=1, flags=re.I)
        if count and normalise_for_comparison(rewritten) != normalise_for_comparison(rewrite_value):
            return f"{bullet_prefix or '•'} {rewritten}" if bullet_prefix else rewritten

    if section.lower() == "summary":
        for pattern, replacement in (
            (r"^\s*i am an?\b", "Professional"),
            (r"^\s*i have\b", "Brings"),
            (r"^\s*i enjoy\b", "Focuses on"),
            (r"^\s*i work\b", "Works"),
        ):
            rewritten, count = re.subn(pattern, replacement, rewrite_value, count=1, flags=re.I)
            if count and normalise_for_comparison(rewritten) != normalise_for_comparison(rewrite_value):
                return rewritten

    if section.lower() in {"skills", "education"}:
        # This is deliberately formatting-only: it cannot add a skill or credential.
        rewritten = re.sub(r"\s*,\s*", " · ", rewrite_value)
        if rewritten != rewrite_value:
            return rewritten

    # A bullet marker is a safe, section-appropriate improvement when the wording
    # contains no weak phrase to rewrite.
    if section.lower() in {"experience", "projects", "experience / projects"}:
        if bullet_prefix != "•":
            return "• " + rewrite_value
    # Keep the fallback distinct without adding a claim: this labels the existing
    # text for the user to review rather than fabricating an outcome.
    return f"{section.title()}: {value}"


def _improvements(text: str, target_role: str, user_feedback: str = "", mode: str = "") -> list[dict[str, str]]:
    sections = detect_sections(text)
    suggestions: list[dict[str, str]] = []
    bullets = _bullet_suggestions(text)
    suggestions.extend(bullets[:3])
    if not sections["summary"] and text.strip():
        first_lines = _content_lines(text)[:3]
        current = "\n".join(first_lines)
        suggestions.append({
            "section": "Summary",
            "current": current,
            "suggested": current + (
                f"\nTarget role focus: {target_role}." if target_role else
                "\nAdd a concise summary of your direction and strongest evidence."
            ),
            "why": "A summary gives the reader context before the detailed evidence. Review the added wording so it remains factual.",
        })
    if sections["summary"]:
        current = _section_content(text, "summary")
        if current:
            suggestions.append({
                "section": "Summary",
                "current": current,
                "suggested": _deterministic_rewrite(current, "Summary"),
                "why": "Keep the summary focused on the target role and evidence already present; remove broad claims you cannot support.",
            })
    for section in ("experience", "projects"):
        if sections[section] and not any(item["section"] == section.title() for item in suggestions):
            content_lines = _bullet_lines(text, section) or _content_lines(text, section)
            if content_lines:
                current = content_lines[0]
                suggested = _deterministic_rewrite(current, section.title())
                suggestions.append({
                    "section": section.title(),
                    "current": current,
                    "suggested": suggested,
                    "why": "Improves scanability and wording while preserving the supplied section content.",
                })
    if not sections["projects"]:
        suggestions.append({
            "section": "Projects",
            "current": "[projects section is missing]",
            "suggested": "Add a Projects section only for real work. For each project: problem · contribution · technologies · verified result · link.",
            "why": "Project evidence can make technical ability concrete without inventing a project.",
        })
    if not sections["contact"]:
        suggestions.append({
            "section": "Contact",
            "current": "[contact information is not detectable]",
            "suggested": "Add your professional email, phone number, location (optional), and relevant links.",
            "why": "A recruiter needs a clear way to reach you. Add your real details before accepting this template.",
        })
    if mode in {"ats", "ats-optimize", "target-role", "skills", "skills-focus"} and target_role:
        keywords = _keyword_analysis(text, target_role)["missing_keywords"][:4]
        if keywords:
            suggestions.append({
                "section": "Keywords",
                "current": "Target-role terms are not clearly represented.",
                "suggested": "Review these terms against your real experience: " + ", ".join(keywords) + ".",
                "why": "Makes the ATS-style gap visible without asking you to claim skills you do not have.",
            })
    if mode in {"recruiter", "recruiter-optimize"} and text.strip():
        first_line = _content_lines(text)[0]
        suggestions.append({
            "section": "Recruiter focus",
            "current": first_line,
            "suggested": first_line + "\n[Lead with role, contribution, and a verified outcome.]",
            "why": "Lead with the clearest role-relevant evidence, then support it with specific work and verified outcomes.",
        })
    if mode in {"concise", "one-page"} or "one page" in user_feedback.lower():
        long_lines = [line for line in _content_lines(text) if len(line.split()) > 35]
        if long_lines:
            suggestions.append({
                "section": "Conciseness",
                "current": long_lines[0],
                "suggested": "Shorten this bullet to action + task + technology + verified result.",
                "why": "Shorter bullets are easier to scan and support a one-page target without deleting facts automatically.",
            })
    if mode in {"projects", "project-focus"} and sections["projects"]:
        content = _section_content(text, "projects")
        suggestions.append({
            "section": "Projects",
            "current": content,
            "suggested": content + "\n[Place the strongest real project first and add verified impact.]",
            "why": "Bring the strongest project first and make its purpose, technical depth, and verified impact explicit.",
        })
    if user_feedback.strip():
        note = f" User instruction applied: {user_feedback.strip()[:160]}"
        suggestions = [{**item, "why": (item.get("why", "") + note).strip()} for item in suggestions]
    return suggestions[:8]


def analyze_resume_text(text: str, target_role: str = "") -> dict[str, Any]:
    text = (text or "").strip()
    sections = detect_sections(text)
    categories = _category_scores(text, sections, target_role)
    keywords = _keyword_analysis(text, target_role)
    bullets = _bullet_suggestions(text)
    ats = _ats_analysis(text, sections, target_role)
    alignment = _role_alignment(text, target_role)
    projects = _project_analysis(text)
    experience = _experience_analysis(text, target_role)
    strengths: list[str] = []
    weaknesses: list[dict[str, str]] = []
    low = text.lower()
    if sections["skills"]:
        strengths.append("A clearly labeled skills or technologies section is present.")
    if sections["experience"] and any(re.search(rf"\b{re.escape(v)}\b", low) for v in ACTION_VERBS):
        strengths.append("Experience includes action-oriented language.")
    if sections["projects"]:
        strengths.append(f"{len(projects)} project signal(s) were detected for review.")
    if re.search(r"\b\d+(?:\.\d+)?%|\b\d+\+", low):
        strengths.append("The resume includes at least one quantified claim to validate and highlight.")
    if target_role and alignment["matches"]:
        strengths.append("Some target-role terminology is already represented.")
    if not strengths:
        strengths.append("The uploaded text provides a starting point for a more focused rewrite.")
    if not sections["contact"]:
        weaknesses.append({"priority": "High", "item": "Add clearly detectable contact information."})
    if not sections["experience"]:
        weaknesses.append({"priority": "High", "item": "Clarify experience or internship evidence if you have it."})
    if bullets:
        weaknesses.append({"priority": "High", "item": "Rewrite generic responsibility bullets into action + task + technology + result."})
    if keywords["missing_keywords"] and target_role:
        weaknesses.append({"priority": "Medium", "item": f"Check missing target-role terms: {', '.join(keywords['missing_keywords'][:4])}."})
    if not re.search(r"\b\d+(?:\.\d+)?%|\b\d+\+", low):
        weaknesses.append({"priority": "Medium", "item": "Add measurable achievements only where you can verify the numbers."})
    if len(text.split()) > 1100:
        weaknesses.append({"priority": "Low", "item": "Reduce long sections to the most relevant evidence."})
    while len(weaknesses) < 3:
        weaknesses.append({"priority": "Low", "item": "Make section headings, dates, and technology context consistent."})
    score = _format_score(sum(categories.values()) / len(categories))
    action_plan = [
        item["item"] for item in sorted(
            weaknesses,
            key=lambda item: {"High": 0, "Medium": 1, "Low": 2}.get(item["priority"], 3),
        )
    ][:5]
    action_plan.extend([
        "Tailor the summary and strongest evidence to the target role.",
        "Keep every suggested technology and result grounded in your actual experience.",
    ])
    return {
        "score": score,
        "category_scores": categories,
        "categories": categories,
        "word_count": len(text.split()),
        "character_count": len(text),
        "sections": sections,
        "missing_sections": [name for name, present in sections.items() if not present and name != "contact"],
        "original_text": text,
        "text": text,
        "target_role": target_role,
        "ats_analysis": ats,
        "role_alignment": alignment,
        "keywords": keywords,
        "bullet_suggestions": bullets,
        "project_analysis": projects,
        "experience_analysis": experience,
        "recruiter_review": {
            "first_impression": "The resume can be assessed from the extracted text; prioritise the clearest evidence first.",
            "strongest_signal": strengths[0],
            "biggest_concern": weaknesses[0]["item"],
            "missing_evidence": weaknesses[1]["item"] if len(weaknesses) > 1 else "Add role-relevant evidence where truthful.",
            "most_important_change": action_plan[0],
            "would_continue_reading": "Potentially, if the strongest evidence is made easier to scan.",
            "disclaimer": "AI-generated recruiter-style simulation. This is not an actual recruiter's opinion or hiring decision.",
        },
        "strengths": strengths[:6],
        "weaknesses": weaknesses[:8],
        "action_plan": action_plan[:5],
        "recommendations": action_plan[:5],
        "improvements": _improvements(text, target_role),
        "analysis_mode": "rule-based",
        "mode": "rule-based",
        "analysis_label": "Rule-based analysis",
    }


def build_improvement_fallback(
    resume_text: str,
    target_role: str = "",
    user_feedback: str = "",
    mode: str = "",
) -> dict[str, Any]:
    suggestions = _improvements(resume_text.strip(), target_role.strip(), user_feedback, mode)
    suggestions = [
        item for item in suggestions
        if not validate_improvement(
            resume_text, item.get("current", ""), item.get("suggested", "")
        )
    ]
    return {
        "suggestions": suggestions,
        "instruction": user_feedback.strip(),
        "mode": "rule-based",
        "analysis_mode": "rule-based",
        "analysis_label": "Rule-based suggestions",
    }